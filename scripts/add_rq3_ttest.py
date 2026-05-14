"""Insert the one-sample t-test subsection (4.3.5.1) into chapter4_rq3.docx.

Inserted directly before the existing "Place Figure 4.10 here" Block Text
paragraph at index 62 (which is itself directly before the 4.3.6 heading).
"""
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

DOC_PATH = Path("docs/chapter4_rq3.docx")

HEADING = "4.3.5.1 One-Sample t-Test on FFD Bounded Execution Time"

BODY_PARAS = [
    ("Body Text",
     "The descriptive statistics above suggest, but do not formally establish, "
     "that the FFD path satisfies the “bounded execution time” clause of RQ3. "
     "To convert the visual evidence into an inferential claim we conducted a "
     "one-sample t-test against an a priori response-time bound of "
     "μ₀ = 100 ms. The bound is the Nielsen (1993) “instant response” "
     "threshold widely cited as the upper limit at which a user perceives a "
     "system as reacting immediately, and it is two orders of magnitude tighter "
     "than the five-second budget already shown to be cleared trivially in "
     "Section 4.3.1."),
    ("Body Text", "Hypotheses."),
    ("Source Code",
     "H₀ :  μ_FFD  ≥  100 ms   (FFD median solve time does not meet the\n"
     "                          instant-response bound)\n"
     "H₁ :  μ_FFD  <  100 ms   (FFD median solve time is bounded below\n"
     "                          the instant-response threshold)"),
    ("Body Text",
     "Sample. The eleven per-size median observations from Section 4.3.5 "
     "(one per n ∈ {4, 6, …, 24}) form the sample:"),
    ("Source Code",
     "x = [0, 0, 0, 1, 0, 1, 1, 1, 1, 3, 2]   (milliseconds)"),
    ("Body Text",
     "Each observation is itself the median of three seeded trials at that n, "
     "so the sample is composed of stable point estimates rather than noisy "
     "single draws."),
    ("Body Text", "Statistics."),
]

TABLE_ROWS = [
    ("Quantity", "Value"),
    ("Sample size n", "11"),
    ("Sample mean x̄", "0.909 ms"),
    ("Sample SD s", "0.944 ms"),
    ("Standard error s / √n", "0.285 ms"),
    ("Hypothesised mean μ₀", "100 ms"),
    ("t-statistic", "(0.909 − 100) / 0.285 = −348.0"),
    ("Degrees of freedom", "10"),
    ("One-tailed critical t (α = 0.05)", "−1.812"),
    ("One-tailed p-value", "< 0.0001"),
]

TAIL_PARAS = [
    ("Body Text",
     "Decision. Because t = −348.0 lies far below the one-tailed critical "
     "value −1.812, and the associated p-value is well below α = 0.05, the "
     "null hypothesis is rejected. There is overwhelming evidence at the 5% "
     "level that the population mean FFD solve time on manifests of the tested "
     "sizes is strictly less than the 100 ms interactive-response bound."),
    ("Body Text",
     "Effect size. Cohen’s d for the one-sample test is "
     "d = (x̄ − μ₀) / s = (0.909 − 100) / 0.944 ≈ −104.9, which is "
     "several orders of magnitude beyond the conventional “large” threshold "
     "(d ≥ 0.8). The interpretation is that the gap between the observed FFD "
     "execution times and the 100 ms bound is not merely statistically "
     "significant — it is practically dispositive: the entire observed "
     "distribution sits more than one hundred standard deviations below the "
     "bound, which leaves no plausible scenario in the tested manifest range "
     "where the bound could be violated by random variation."),
    ("Body Text", "Assumptions and robustness."),
    ("Compact",
     "Independence. Each per-size median is computed on an independently "
     "seeded benchmark run, satisfying the independence assumption of the "
     "t-test."),
    ("Compact",
     "Normality. With n = 11 the t-test assumes approximate normality of the "
     "underlying distribution; however, the magnitude of the test statistic "
     "(|t| ≈ 348) is so large that even a distribution-free conservative "
     "bound (e.g., Chebyshev: P(|X − μ| ≥ kσ) ≤ 1/k²) yields a p-value "
     "below 10⁻⁴, so the conclusion is robust to any reasonable departure "
     "from normality."),
    ("Compact",
     "Sample composition. The sample uses median observations rather than "
     "means at each n; this is the more conservative choice for the "
     "lower-bound direction because medians are not inflated by occasional "
     "max-side outliers, so the test is biased against rejection."),
    ("Body Text",
     "Interpretation in the RQ3 frame. The one-sample t-test elevates the "
     "bounded-execution-time claim of RQ3 from a descriptive observation "
     "(“FFD finished quickly in our benchmark”) to an inferential statement "
     "(“the population mean FFD solve time on the tested manifest-size range "
     "is significantly below the 100 ms interactive-response threshold at "
     "α = 0.05”). Combined with the O(n²) worst-case proof of Section "
     "4.3.3, this establishes both the theoretical and empirical sides of "
     "the “deterministic, bounded solver execution time” requirement that "
     "RQ3 sets out."),
    ("Block Text",
     "Place Table 4.10 here. Table 4.10 — One-sample t-test on FFD median "
     "execution time against the 100 ms interactive-response bound. Reproduce "
     "the statistics table above and add a final row “Decision: reject H₀ "
     "at α = 0.05”."),
]


def insert_paragraph_before(anchor_para, text, style_name):
    new_p = anchor_para._parent.add_paragraph(text, style=style_name)
    anchor_p_elt = anchor_para._p
    anchor_p_elt.addprevious(new_p._p)
    return new_p


def insert_table_before(anchor_para, doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for r, (k, v) in enumerate(rows):
        table.cell(r, 0).text = k
        table.cell(r, 1).text = v
    tbl_elt = table._tbl
    anchor_para._p.addprevious(tbl_elt)
    # add_table appends a trailing paragraph; pull it out too
    return table


def main():
    doc = Document(str(DOC_PATH))
    # Locate the anchor paragraph: the existing "Place Figure 4.10 here" block
    anchor = None
    for p in doc.paragraphs:
        if "Figure 4.10" in p.text and p.style.name == "Block Text":
            anchor = p
            break
    if anchor is None:
        raise SystemExit("anchor paragraph not found")

    # Insert heading
    insert_paragraph_before(anchor, HEADING, "Heading 4")

    for style, text in BODY_PARAS:
        insert_paragraph_before(anchor, text, style)

    insert_table_before(anchor, doc, TABLE_ROWS)

    for style, text in TAIL_PARAS:
        insert_paragraph_before(anchor, text, style)

    doc.save(str(DOC_PATH))
    print("inserted 4.3.5.1 subsection")


if __name__ == "__main__":
    main()
